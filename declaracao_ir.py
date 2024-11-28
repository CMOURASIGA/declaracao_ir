import pyodbc
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests

# Função para definir a conexão com base no ambiente escolhido
def definir_conexao(ambiente):
    if ambiente == 'HML':
        conn_str = 'DRIVER={SQL Server};SERVER=45.6.154.117,1861;DATABASE=CZ519L_135184_RM_DV;UID=CLT135184GustavoBarros;PWD=qlpkf51796KWRXS@!'
    elif ambiente == 'PRD':
        conn_str = 'DRIVER={SQL Server};SERVER=45.6.154.118,38000;DATABASE=CZ519L_135175_RM_PD;UID=CLT135175userti;PWD=jxlap28517DFBZQ@!'
    else:
        raise ValueError("Ambiente inválido. Escolha HML ou PRD.")
    
    conn = pyodbc.connect(conn_str)
    return conn

# Função para baixar o logo
def baixar_logo(url, nome_arquivo):
    response = requests.get(url)
    if response.status_code == 200:
        with open(nome_arquivo, 'wb') as file:
            file.write(response.content)
        print(f"Logo baixado com sucesso: {nome_arquivo}")
    else:
        print(f"Erro ao baixar logo: {response.status_code}")
        nome_arquivo = None
    return nome_arquivo

# Função para executar a consulta SQL
def executar_consulta(conexao, codcoligada, codfilial, periodo_letivo, ra):
    query = f"""
    SELECT
        SP.RA AS RA,
        SS.NOME AS NOMESERVICO,
        FL.DATAVENCIMENTO AS DATAAUX,
        CONVERT(varchar(30), FL.DATAVENCIMENTO, 103) AS DATAVENCIMENTO,
        CONVERT(decimal(10, 2), TM.VALORBRUTO) AS VALOR,
        SPL.CODPERLET AS ANOLETIVO,
        FC.NOME AS NOMERESP,
        PP.NOME AS NOMEALUNO,
        FC.CGCCFO AS CNPJRESP,
        GFILIAL.NOME AS NOMEFILIAL,
        GFILIAL.RUA AS RUAFILIAL,
        GFILIAL.BAIRRO AS BAIRROFILIAL,
        GFILIAL.CIDADE AS CIDADEFILIAL,
        GFILIAL.COMPLEMENTO AS COMPLEMENTOFILIAL,
        GFILIAL.NUMERO AS NUMEROFILIAL,
        GFILIAL.ESTADO AS ESTADOFILIAL,
        GFILIAL.CGC AS CNPJFILIAL
    FROM SLANMOV SMOV
    INNER JOIN TMOV TM ON SMOV.CODCOLIGADA = TM.CODCOLIGADA AND SMOV.IDMOV = TM.IDMOV
    INNER JOIN FCFO FC ON FC.CODCFO = TM.CODCFO
    LEFT JOIN SLAN SL ON SL.CODCOLIGADA = SMOV.CODCOLIGADA AND SL.IDLAN = SMOV.IDLAN
    INNER JOIN FLAN FL ON FL.CODCOLIGADA = SL.CODCOLIGADA AND FL.IDLAN = SL.IDLAN
    LEFT JOIN SPARCELA SP ON SP.CODCOLIGADA = SL.CODCOLIGADA AND SP.IDPARCELA = SL.IDPARCELA
    INNER JOIN SALUNO SA ON SA.RA = SP.RA AND SA.CODCOLIGADA = SP.CODCOLIGADA
    INNER JOIN PPESSOA PP ON PP.CODIGO = SA.CODPESSOA
    LEFT JOIN SPLETIVO SPL ON SP.CODCOLIGADA = SPL.CODCOLIGADA AND SP.IDPERLET = SPL.IDPERLET
    LEFT JOIN SSERVICO SS ON SP.CODCOLIGADA = SS.CODCOLIGADA AND SP.CODSERVICO = SS.CODSERVICO
    INNER JOIN GFILIAL GFILIAL ON GFILIAL.CODCOLIGADA = {codcoligada} AND GFILIAL.CODFILIAL = {codfilial}
    WHERE TM.STATUS = 'N'
      AND FL.STATUSLAN IN (1, 4)
      AND SPL.CODPERLET = '{periodo_letivo}'
      AND SP.RA = '{ra}'
      AND FL.CODTDO NOT IN ('0003', '0006', '0355')
    ORDER BY DATAAUX ASC
    """
    cursor = conexao.cursor()
    cursor.execute(query)
    resultados = cursor.fetchall()
    colunas = [desc[0] for desc in cursor.description]
    return resultados, colunas

# Função para tratar valores
def tratar_valor(valor):
    if valor is None:
        return ""
    return str(valor)

# Função para gerar o relatório
def gerar_relatorio(resultados, colunas, nome_arquivo, responsavel, cpf, aluno, ano, cidade='Brasília', logo_path=None):
    document = Document()

    # Ajustar margens do documento
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.8)

    # Adicionar logo no cabeçalho
    if logo_path:
        header = document.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.5))  # Tamanho ajustável do logo

    # Cabeçalho e introdução
    nome_filial = resultados[0][colunas.index('NOMEFILIAL')]
    document.add_heading(nome_filial, level=1)
    document.add_paragraph('Declaração de quitação de débitos', style='Title')
    document.add_paragraph(
        f"Pelo presente instrumento, {nome_filial}, estabelecimento inscrito no CNPJ sob nº "
        f"{resultados[0][colunas.index('CNPJFILIAL')]}, situado na {resultados[0][colunas.index('RUAFILIAL')]}, "
        f"nº {resultados[0][colunas.index('NUMEROFILIAL')]}."
    )
    # Tabela
    table = document.add_table(rows=1, cols=5)
    headers = ['Parcela', 'Data Vencimento', 'Serviço', 'Data Pagamento', 'Valor Pago']
    table.style = 'Table Grid'

    # Definir larguras uniformes para as colunas
    widths = [Inches(0.5), Inches(1.0), Inches(2.5), Inches(1.0), Inches(1.0)]  # Ajuste as larguras conforme necessário

    # Preencher os cabeçalhos
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Centralizar os títulos dos cabeçalhos
    header_cells = table.rows[0].cells
    for idx, cell in enumerate(header_cells):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Ajustar a largura dos cabeçalhos
        cell.width = widths[idx]

    # Preencher os dados
    for i, row in enumerate(resultados):
        cells = table.add_row().cells
        cells[0].text = str(i + 1)
        cells[1].text = tratar_valor(row[colunas.index('DATAVENCIMENTO')])
        cells[2].text = tratar_valor(row[colunas.index('NOMESERVICO')])
        
        # Ajuste para a data de pagamento
        if isinstance(row[colunas.index('DATAAUX')], datetime):
            cells[3].text = row[colunas.index('DATAAUX')].strftime("%d/%m/%Y")
        else:
            cells[3].text = datetime.strptime(row[colunas.index('DATAAUX')], "%Y-%m-%d").strftime("%d/%m/%Y")
        
        # Ajuste para o valor
        cells[4].text = f"R$ {row[colunas.index('VALOR')]:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        # Aplicar largura uniforme nas células
        for idx, cell in enumerate(cells):
            cell.width = widths[idx]

    # Texto final adicionado uma única vez
    document.add_paragraph(
        f"Vinculadas ao período letivo {ano} e referente ao Contrato de Prestação de Serviços Educacionais celebrado com o Sr(a). {responsavel}, "
        f"inscrito(a) no CPF sob nº {cpf}, para o aluno {aluno}."
    )
    data_atual = datetime.now().strftime('%d de %B de %Y')
    document.add_paragraph(f"{cidade}, {data_atual}.")

    # Salvando o documento
    document.save(nome_arquivo)
    print(f"Relatório salvo em {nome_arquivo}")

# Execução principal
if __name__ == '__main__':
    link_logo = "https://drive.google.com/uc?export=download&id=1dZU6WNqQJXyNiaNdd9CeusIFYBXRzo-b"
    caminho_logo = "logo_empresa.png"
    logo_path = baixar_logo(link_logo, caminho_logo)

    ambiente = input("Digite o ambiente (HML ou PRD): ").strip()
    codcoligada = int(input("Digite o código da coligada: ").strip())
    codfilial = int(input("Digite o código da filial: ").strip())
    periodo_letivo = input("Digite o período letivo (exemplo: 2024): ").strip()
    ra = input("Digite o RA do aluno: ").strip()
    caminho_salvar = input("Digite o caminho para salvar o relatório: ").strip()
    nome_arquivo = f"{caminho_salvar}/Declaracao_Quitacao_{ra}.docx"

    conexao = definir_conexao(ambiente)
    resultados, colunas = executar_consulta(conexao, codcoligada, codfilial, periodo_letivo, ra)

    responsavel = resultados[0][colunas.index('NOMERESP')]
    cpf = resultados[0][colunas.index('CNPJRESP')]
    aluno = resultados[0][colunas.index('NOMEALUNO')]

    gerar_relatorio(resultados, colunas, nome_arquivo, responsavel, cpf, aluno, periodo_letivo, logo_path=logo_path)
